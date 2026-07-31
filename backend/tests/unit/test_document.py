"""文档提取单测。"""

from __future__ import annotations

import csv as _csv
import io
import json

import pytest

from xuwen.chat_api.document import (
    DocumentError,
    extract,
    supported_extensions,
)


def test_supported_extensions_includes_common_formats():
    exts = supported_extensions()
    for e in ["txt", "md", "json", "csv", "pdf", "docx", "xlsx", "html"]:
        assert e in exts


def test_extract_txt():
    data = "你好\n世界".encode()
    doc = extract(data, "hello.txt")
    assert doc.extension == "txt"
    assert "你好" in doc.text
    assert doc.char_count >= 4
    assert doc.estimated_tokens > 0


def test_extract_md():
    data = "# 标题\n\n正文一段".encode()
    doc = extract(data, "note.md")
    assert "# 标题" in doc.text


def test_extract_json_pretty():
    raw = b'{"a": 1, "b": "\xe4\xbd\xa0\xe5\xa5\xbd"}'
    doc = extract(raw, "data.json")
    parsed = json.loads(doc.text)
    assert parsed["a"] == 1
    assert parsed["b"] == "你好"


def test_extract_csv_to_markdown_table():
    buf = io.StringIO()
    writer = _csv.writer(buf)
    writer.writerow(["name", "age"])
    writer.writerow(["示例甲", "8"])
    writer.writerow(["示例乙", "9"])
    data = buf.getvalue().encode("utf-8")
    doc = extract(data, "students.csv")
    assert "| name | age |" in doc.text
    assert "| 示例甲 | 8 |" in doc.text


def test_extract_csv_truncates_large():
    buf = io.StringIO()
    writer = _csv.writer(buf)
    writer.writerow(["i"])
    for i in range(300):
        writer.writerow([str(i)])
    doc = extract(buf.getvalue().encode(), "big.csv")
    assert "仅展示前 200 行" in doc.text


def test_extract_html_strips_scripts():
    data = """
    <html>
    <head><title>X</title></head>
    <body>
      <h1>大标题</h1>
      <script>var x = 1;</script>
      <p>正文段落</p>
    </body>
    </html>
    """.encode()
    doc = extract(data, "page.html")
    assert "大标题" in doc.text
    assert "正文段落" in doc.text
    assert "var x = 1" not in doc.text


def test_extract_pdf_minimal():
    """空 PDF 提取后内容为空，应抛 DocumentError。"""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf 未安装")

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)

    with pytest.raises(DocumentError):
        extract(out.getvalue(), "blank.pdf")


def test_extract_docx_basic():
    """生成一份简单 DOCX 测试提取。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    result = extract(out.getvalue(), "test.docx")
    assert "第一段" in result.text
    assert "第二段" in result.text


def test_extract_xlsx_basic():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["姓名", "分数"])
    ws.append(["示例甲", 88])
    ws.append(["示例乙", 92])
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    result = extract(out.getvalue(), "scores.xlsx")
    assert "姓名" in result.text
    assert "示例甲" in result.text
    assert "92" in result.text


def test_extract_unknown_extension_rejected():
    with pytest.raises(DocumentError) as exc:
        extract(b"binary", "mystery.bin")
    assert "暂不支持" in exc.value.message


def test_extract_oversize_rejected():
    with pytest.raises(DocumentError) as exc:
        extract(b"x" * 10, "tiny.txt", max_bytes=5)
    assert "过大" in exc.value.message


def test_extract_empty_after_parsing_rejected():
    with pytest.raises(DocumentError):
        extract(b"   \n   ", "blank.txt")


def test_extract_invalid_json_rejected():
    with pytest.raises(DocumentError):
        extract(b"{ not json }", "bad.json")


def test_extract_chinese_decoding_fallback():
    """非 UTF-8 编码的中文 txt 也应能解码。"""
    text = "你好世界 GBK"
    data = text.encode("gbk")
    doc = extract(data, "gbk.txt")
    assert "你好" in doc.text
